#! /usr/bin/env python

import string
import os
import os.path
import sys
import math
import random
from docx import Document
from docx.shared import Cm
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH



######################################################################################################
#                                                                                                    #
# function for reading gaussian output files                                                         #
#                                                                                                    #
# Reads gaussian output files and returns an array containing the cartesian coordinates (float),     #
# the list of atom symbols (str), the list of atom types (str) and, if found, the level layer (str)  #
#  of each atom from ONIOM calcs.                                                                    #
#                                                                                                    #
# The file passed as argument must include the ".out" extension.                                     # 
# The atom symbols, types and connectivities are read from the begining of the file.                 # 
# The cartesian coordinates are updated for every step in the case of a geometry optimization.       #
# In case of final frequency calculation after an optimization, or if the file only contains         #
# a frequency calculation, the cartesian coordinates are read from this section.                     #                          
######################################################################################################

def read_cartsandlayer_from_g_out(file):

    outfile=open(file,"r")
    outlines=outfile.readlines()

    i=0
    #read initial geometry; layers, atom symbols and atom types are read here as well
    while i<len(outlines):

        if outlines[i].strip().find("Symbolic Z-matrix:")>-1:
            i=i+1
            while outlines[i].strip().startswith("Charge"):
                i=i+1
            molecule=[]
            cartesians=[]
            symbols=[]
            types=[]
            layers=[]
            while outlines[i].strip(' ')!="\n":
                if outlines[i].split()[0].find("-")>0:
                    atom_symbol=outlines[i].split()[0].split("-")[0]
                    atom_type=outlines[i].split()[0].split("-")[1]
                else: 
                    atom_symbol=outlines[i].split()[0]
                    atom_type=""
 
                if len(outlines[i].split())>4 and len(outlines[i].split()[1])<3:
                    atom_x=float(outlines[i].split()[2])
                    atom_y=float(outlines[i].split()[3])
                    atom_z=float(outlines[i].split()[4])
                else:
                    atom_x=float(outlines[i].split()[1])
                    atom_y=float(outlines[i].split()[2])
                    atom_z=float(outlines[i].split()[3])

                if len(outlines[i].split())>5:
                    atom_layer=outlines[i].split()[5]
                elif len(outlines[i].split())>4 and len(outlines[i].split[1])>3:
                    atom_layer=outlines[i].split()[4]
                else:
                    atom_layer=""
            
                symbols.append(atom_symbol)
                types.append(atom_type)
                cartesians.append([atom_x,atom_y,atom_z])
                layers.append(atom_layer)
                i=i+1
            

        #update geometry for every step
        if (outlines[i].strip().find("GradGradGradGradGrad")>-1) and (outlines[i-1].strip().find("Predicted")>-1):

            new_step=False

            j=0
            while j<20:
                if outlines[i].find("orientation")>1 and outlines[i+2].find("Coordinates")>1:
                    new_step=True
                    break
                i=i+1
                j=j+1

            if new_step:
                i=i+5
                cartesians=[]  
                while not (outlines[i].strip(' ').startswith("-------------------------------------------------------")):
                    if len(outlines[i].split())>1:
                        atom_x=float(outlines[i].split()[3])
                        atom_y=float(outlines[i].split()[4])
                        atom_z=float(outlines[i].split()[5])
                        cartesians.append([atom_x,atom_y,atom_z])
                    i=i+1


        #update geometry if a link1 hessian job is at the end
        if (outlines[i].find("Redundant internal coordinates found in file.  (old form).")>-1):
            i=i+1
            cartesians=[]      
            while outlines[i].find("Recover connectivity data from disk.")==-1:
                if len(outlines[i].split(","))>1:  
                    atom_x=float(outlines[i].split(",")[2])
                    atom_y=float(outlines[i].split(",")[3])
                    atom_z=float(outlines[i].split(",")[4])
                    cartesians.append([atom_x,atom_y,atom_z])                
                i=i+1
        
        i=i+1

    outfile.close()
    return [cartesians,symbols,types,layers]
#############################End of functions for reading coordinates#################################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############

######################################################################################################
#                                                                                                    #
# Function for reading the energy of a gaussian output file. Returns a string including the level    #
# of theory and the units. In the case of ONIOM calculations, returns the extrapolated energy.       #
# The file passed as argument must include the ".out" extension.                                     #
#                                                                                                    #
######################################################################################################

def read_text_energy(file):

    outfile=open(file,"r")
    outlines=outfile.readlines()
    energy=""
    level_of_theory=""

    i=0
    while i<len(outlines):
        if outlines[i].find("SCF Done:")>-1:
            level_of_theory=outlines[i].split()[2]
            energy=outlines[i].split()[4]
        
        if outlines[i].find(" ONIOM: extrapolated energy = ")>-1:
            level_of_theory="ONIOM: extrapolated energy"
            energy=outlines[i].split()[4]
        i=i+1
    outfile.close()
    return level_of_theory+":\n"+energy+" Hartrees"
##########################End of functions for reading energy information#############################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############

######################################################################################################
#                                                                                                    #
# Function for reading the first three imaginary frequencies found in a gaussian output file.        #
# Returns "No imaginary frequencies found" if there are no imaginary frequencies, "Single imaginary  #
# freq: "+value if there is only one, or "imaginary freqs:"+ a list of 2 or 3 freqs if there are     #
# than one (does not distinguish if there are more than 3 img. freqs).                               #
# The file passed as argument must include the ".out" extension.                                     #
#                                                                                                    #
######################################################################################################

def read_img_freqs(file):

    outfile=open(file,"r")
    outlines=outfile.readlines()
    i=0
    s=""
    while i<len(outlines):
        
        if outlines[i].strip().find("and normal coordinates:")>-1:   
            s=""
            i=i+3
            j=2
            while j<5:
                if float(outlines[i].split()[j])<0: 
                    s=s+" "+outlines[i].split()[j]
                j=j+1          
        i=i+1
    if s=="": s="No imaginary frequencies found"
    elif len(s.split())==1: s= "Single imaginary freq: "+s
    elif len(s.split())>1: s= str(len(s.split()))+" imaginary freqs: "+s
    outfile.close()
    return s
########################End of functions for reading frequency information############################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# Functions for creating .pdb files from molecules.                                                  #
# The atoms are defined as HETAM but the residue id. is stablished using the ONIOM layer             #
# Connectivity is not set.                                                                           #
# than one (does not distinguish if there are more than 3 img. freqs).                               #
# The file passed as argument must include the ".out" extension.                                     #
#                                                                                                    #
######################################################################################################

def text_in_pdb (cartesians,symbols,types,layers):

    i=0
    text="COMPND    UNNAMED\nAUTHOR yo\n"

    while i<len(symbols):
         
        if layers[i]=="H" or layers[i]=="": resid="unk1"
        elif layers[i]=="L": resid="unk2"
        elif layers[i]=="M": resid="unk2"
        text=text+"HETATM%5i  %-4s%-4s" %(i+1,symbols[i],resid)
        text=text+"    1     %- 8.3f%- 8.3f%- 8.3f" %(cartesians[i][0],cartesians[i][1],cartesians[i][2])
        text=text+" 1.00  0.00           "+symbols[i]+"\n"

        i=i+1
    return text

def molecule_to_text_in_pdb (molecule):
    return text_in_pdb(molecule[0],molecule[1],molecule[2],molecule[3])

# the function that should be called:
def molecule_to_pdb (molecule,filename):
    pdbfile= open(filename,"w")
    pdbfile.write(molecule_to_text_in_pdb(molecule))
    pdbfile.close()
#############################End of functions for writting pdf files##################################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# functions for rotating a molecule with euler angles. Perform a XYZ rotation                        #
#                                                                                                    #
######################################################################################################

def rotate_vector(vector,alpha,beta,gamma):

    c1= math.cos(alpha)
    s1= math.sin(alpha)
    c2= math.cos(beta)
    s2= math.sin(beta)
    c3= math.cos(gamma)
    s3= math.sin(gamma)

    x=c2*c3*vector[0] - c2*s3*vector[1] + s2*vector[2]
    y=(c1*s3+c3*s1*s2)*vector[0] + (c1*c3-s1*s2*s3)*vector[1] + (-c2*s1)*vector[2]
    z=(s1*s3-c1*c3*s2)*vector[0] + (c3*s1+c1*s2*s3)*vector[1] + (c1*c2)*vector[2]
    return [x,y,z]

def rotate_vectors(vectors,alpha,beta,gamma):

    new_vectors=[]
    for v in vectors:
        new_vectors.append(rotate_vector(v,alpha,beta,gamma))
    return new_vectors

# the function that should be called, takes a molecule array and returns it rotated
def rotate_molecule(molecule,angles):

    new_vectors_in_molecule=rotate_vectors(molecule[0],angles[0],angles[1],angles[2])
    return [new_vectors_in_molecule,molecule[1],molecule[2],molecule[3]]
########################End of functions for rotating cartesian coordinates###########################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


####################NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED####################
def change_cart_in_molecule (newcartesians,molecule):
    return [newcartesians,molecule[1],molecule[2],molecule[3]]


######################################################################################################
#                                                                                                    #
# different alternatives for calculating the "score" funtion for optimizing the molecule orientation #
#                                                                                                    #
######################################################################################################

# sum of the distances of the atoms once they are projected in the XY plane
# minimizing, useful for eclipsing atoms when minimized
def sum_XY_distances_score(vectors):
    i=0
    score=0
    while i<len(vectors):
        j=i+1
        while j<len(vectors):
            dist=(vectors[i][0]-vectors[j][0])**2+(vectors[i][1]-vectors[j][1])**2
            score=score+dist
            j=j+1
        i=i+1
    return score

# sum of the inverse of the distance of the atoms once they are projected in the XY plane
# minimizing, the projection will try to prevent eclipsing atoms
def sum_inv_XY_distances_score(vectors):
    i=0
    score=0.0
    while i<len(vectors):
        j=i+1
        while j<len(vectors):
            dist=(vectors[i][0]-vectors[j][0])**2+(vectors[i][1]-vectors[j][1])**2
            if dist>0.2:
                score=score+(1/dist)**2
            else: score=score+10.0
            j=j+1
        i=i+1
    return score 

# sum of the distance of the atoms to the center after they are projected to to XY plane
# changing sign and minimizing, the projection will occupy larger area; if only two atoms are in vectors
# the "bond" between the two atoms will be longer

def sum_XY_values_score(vectors):
    i=0
    score=0.0
    while i<len(vectors):
        score=score+vectors[i][0]**2+vectors[i][1]**2
        i=i+1
    return score

######################################################################################################
# function that will be called during the optimization                                               #
#                                                                                                    #
# prepare the atoms that will be used in calculating the score according to the                      #
# specifications in the parameter: "method" and calls to the above functions to calculate it         #
#                                                                                                    #
# ***********************************"METHOD" SINTAX*************************************************#
#                                                                                                    #
# prevent_eclipsing         tries to minimize the inverse of the distance of the atoms after         #
#                           projected, affects to all atoms                                          #
#                                                                                                    #
# eclipse_HL                tries to join (eclipse) the atoms, affect to the high level layer if the #
#                           lave information is given                                                #
#                                                                                                    #
# span_notH                 increases the distance to the center after atoms are proyected in the XY #
#                           plane, affect to all atoms but not to H                                  #
#                                                                                                    #
# prevent_eclipsing_HLnotH  affects to non H-atoms in the high level layer                           #
#                                                                                                    #
# prevent_eclipsing_:5-11   tries to prevent eclipsing of atom 5 by 11                               #
#                                                                                                    #
# eclipse_:5-11             searches an orientation in which 5 and 11 are eclipsed                   #
#                                                                                                    #
# span_:5-11,4-8,3-7        searches an orientation in wich the bonds 5-11, 4-8, and 3-7 are         #
#                           shown longer                                                             #
######################################################################################################

def calculate_XY_score(molecule,method):

    score=0
    vectors=[]
    pairs=[]
    this_vector=[]
    if method=="prevent_eclipsing":
        score= sum_inv_XY_distances_score(molecule[0])

    if method.find("_HL")>-1:
  
        i=0
        this_vector=[]
        while i<len(molecule[3]):
            if molecule[3][i]!="L" and molecule[3][i]!="M":
                if (molecule[1][i]!="H") or (molecule[1][i]=="H" and method.find("notH")>-1):
                    this_vector.append(molecule[0][i])
            i=i+1
        vectors=[this_vector]

    if method.find("_notH")>-1:
        i=0
        this_vector=[]
        while i<len(molecule[3]):
            if molecule[1][3]!="H":
                this_vector.append(molecule[0][i])
            i=i+1
        vectors=[this_vector]

    if method.find(":")>-1:
        this_vector=[]
        if len(method.split(":"))>1:
            if len(method.split(":")[1].split(","))>0 :
                for v in method.split(":")[1].split(","):

                    if len(v.split("-"))>1:
                        for vv in v.split("-"):

                            this_vector.append(molecule[0][int(vv)-1])
                vectors.append(this_vector)
            else: print("something failed when trying to read a list of atoms");sys.exit()
        else: print("something failed when trying to read a list of atoms");sys.exit()


    if method.startswith("prevent_eclipsing_"):
        score=0
        for v in vectors:
            score= score+sum_inv_XY_distances_score(v)

    if method.startswith("eclipse_"):
        score=0
        for v in vectors:
            score= score+sum_XY_distances_score(v)

    if method=="span":
        score=-sum_XY_values_score(molecule[0])
    
    if method.startswith("span_"):
        score=0
        for v in vectors:
            score= score+sum_XY_values_score(v)    
    
    return score 


######################################################################################################
#                                                                                                    #
# find the best oientation of a molecule according to the criteria described in "method"             #
#               ************(see METHOD SINTAX above)*************                                   #
# returns a list (conventional python array), nor a numpy array                                      #
######################################################################################################

def get_best_orientation(molecule,method):

    # libraries imported here to prevent errors if this feature is not required and modules are not installed    
    import numpy as np
    from scipy import optimize
    from scipy.optimize import Bounds
    from scipy.optimize import basinhopping

    # the function that will be minimized
    def obj_function(x):
        return calculate_XY_score(rotate_molecule(molecule,x),method) 

    # tries L-BFGS-B grandient based method (with restraints) and if it fails (quite ofter) tries Nelder-Mead   
    """
    bounds = Bounds([0, 0, 0], [2*math.pi, 2*math.pi, 2*math.pi])
    res=optimize.minimize(obj_function,np.array([1,1,1]),method="L-BFGS-B",bounds=bounds)
    if not res.success:
        #for debugging:
        #res=optimize.minimize(obj_function,np.array([1,1,1]),method="SLSQP",bounds=bounds,options={'disp': 1})
        res=optimize.minimize(obj_function,np.array([1,1,1]),method="Nelder-Mead")
    """

    # uses the more robust but slower Nelder-Mead method... more consistent but slower, no restraints
    res=optimize.minimize(obj_function,np.array([1,1,1]),method="Nelder-Mead")
        
    # for debuggin:
    #print res
    return res.x.tolist()

#########################End of functions for obtaining best orientation##############################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# function for returning an string with the cartesian coordinates, including the ONIOM layer         #
# (if present)                                                                                       #
#                                                                                                    #
######################################################################################################

def prepare_cart(molecule):
    cart=""
    i=0
    while i<len(molecule[1]):
        cart=cart+" %2s    %- 10.6f%- 10.6f%- 10.6f %2s \n" %(molecule[1][i],molecule[0][i][0],molecule[0][i][1],molecule[0][i][2],molecule[3][i])
        i=i+1
    return cart
########################End of functions for obtaining cartesian coordinates #########################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# renders a png image using pymol from a pdb file using the settings specified in pml_file           #
# default size is good enough for a 2-column word document                                           #
#                                                                                                    #
######################################################################################################

def render_image(pdb_file,pml_file,xwidth=800,ywidth=800,filename=""):

    if filename=="": png_file=pdb_file.split('.')[0]+".png"

    
    import pymol
    import __main__
    __main__.pymol_argv = [ 'pymol', '-Qc']
    
    pymol.finish_launching()
    pymol.cmd.load(pdb_file,pdb_file.split('.')[0])
    # to prevent filling the shell with lots of messages:
    sys.stdout = open('log.txt', 'w')
    pymol.cmd.run(pml_file,"local")
    pymol.cmd.zoom("all","1")
    pymol.cmd.png(png_file,xwidth,ywidth)
    # no more disturbing messages in the shell, so recover stdout valor:
    sys.stdout= sys.__stdout__
    pymol.cmd.sync()    ###########################################################################
    pymol.cmd.delete("all")
    # if quit() is executed, pymol will not open after the first call to this function (why????)
    #pymol.cmd.quit()
 
###############################End of functions for rendering image ##################################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# function for creating and formatting with 2 colums a word document.                                # 
# returns a docx Document object                                                                     #
#                                                                                                    #
######################################################################################################

def create_and_format_document(title=""):

    doc=Document()
    p1=doc.add_paragraph()
    p1.alignment=WD_ALIGN_PARAGRAPH.CENTER
    t1=p1.add_run(title)
    t1.font.size=Pt(24)
    doc.add_section(WD_SECTION.CONTINUOUS)
    section = doc.sections[-1]
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')
    return doc


######################################################################################################
#                                                                                                    #
# function for writting in the word document. Takes a Document object (should be created before)     # 
# and adds the text in the argument. If given, it also places the picture passed as png              #
#                                                                                                    #
######################################################################################################

def write_in_document(doc,struct_name,energy,freq,cartesians,png_file="",pict_size=7):

    p1=doc.add_paragraph()
    t1=p1.add_run(struct_name)
    t1.font.bold=True
    
    p2=doc.add_paragraph()
    t2=p2.add_run(energy)

    p3=doc.add_paragraph()
    t3=p3.add_run(freq)
    if freq!="No imaginary frequencies found":
        t4=p3.add_run(" cm")
        t5=p3.add_run("-1")
        t5.font.superscript=True      
    if png_file!="":
        doc.add_picture(png_file,width=Cm(pict_size))

    p4=doc.add_paragraph()
    t6=p4.add_run(cartesians)
    t6.font.size=Pt(11)

def save_document(doc,name):
    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.save(name)

######################End of functions for writing results in a word document#########################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# Functions for working with files and directories. They allow to specify all the files that will    #
# be used for getting the frequencies, energies and coordinates simply giving two files:             #
# one of the files (and its route) with frequency information and one of the files with the energy.  #
#                                                                                                    #
# The files have the form:                          prefix_for_freq+NAME+suffix_for_freq.ext         #
#                                                     prefix_for_sp+NAME+suffix_for_sp.ext           #
#                                                                                                    #
#  - file with frequency information and with energy information for a single structure MUST share   #
#  the same NAME. It will be used to name the structures in the final .doc file.                     #
#                                                                                                    #
#  - all files with frequency information MUST share  prefix_for_freq and suffix_for_freq; all files #
#  with energy information MUST share  prefix_for_so and suffix_for_sp.                              #
#                                                                                                    #
#  - all files with frequency information MUST be in the same directory; all files with energy       #
#  information MUST be in the same directory. All files can be in a single directory.                #
#                                                                                                    #
#  - prefix_for_freq, suffix_for_freq, prefix_for_sp, suffix_for_sp could be empty.                  #
#                                                                                                    #
#  - the extension of the output files could be anything (not neccesarily "out" or "log"), so the    #
#  input files (generaly "com" or "inp") should not be present in the same directory.                #
#                                                                                                    #
#  - if the same file (name and route) is specified for frequency and energy information, a single   #
#  file will be used for both                                                                        #
#                                                                                                    #
######################################################################################################

#self explanatory...
def common_in_two_strings(string1,string2):

    common_in_strings=""
    if len(string1)>len(string2):
        short_string=string2
        long_string=string1
    else: 
        short_string=string1
        long_string=string2
    j=0
    while j<len(short_string):
        i=len(short_string[j:])
        while i>0:
            if long_string.find((short_string[j:i]))>-1:
                if len(short_string[j:i])>len(common_in_strings):
                    common_in_strings=short_string[j:i]

                break
            i=i-1
        j=j+1
    return common_in_strings
    
#returns an array with the route to a file on one side and the file in the other
def split_route_and_file(filename):
    return [filename.replace(filename.split("/")[-1],""),filename.split("/")[-1]]

#self explanatory... (the dot is not included)
def remove_file_extension(filename):
    if filename.find(".")>-1:
        return filename.replace(filename.split(".")[-1],"")[:-1]
    else: return filename


####################NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED####################
#same as before, but returns an array with the name of the file (extension removed) and the extension.
def split_file_extension(filename):
    return [filename.replace(filename.split(".")[-1],"")[:-1],filename.split(".")[-1]]



#given a file name (with or without route) and a name, it returns an array with the prefix and a suffix
def prefix_and_suffix(filename,name):
    if (filename.find("/")==-1):
        return remove_file_extension(filename).split(name)
    else:
        return remove_file_extension(split_route_and_file(filename)[1]).split(name)

# from two files, returns a double array with the prefix and suffix used in both of them
def get_prefix_and_suffix(freq_file,sp_file):
    f1=""
    if (freq_file.find("/")==-1): f1=remove_file_extension(freq_file)
    else: f1=remove_file_extension(split_route_and_file(freq_file)[1])
    f2=""
    if (sp_file.find("/")==-1): f2=remove_file_extension(sp_file)
    else: f2=remove_file_extension(split_route_and_file(sp_file)[1])   
    name=common_in_two_strings(f1,f2)
    freq_prefix_and_suffix=prefix_and_suffix(f1,name)
    sp_prefix_and_suffix=prefix_and_suffix(f2,name)

    return [freq_prefix_and_suffix,sp_prefix_and_suffix]


####################NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED_NOT_USED####################
#return the files in a directory that start with prefix and end with suffix.
def list_files_with_prefix_and_suffix(route,prefix,suffix):
    all_files=os.listdir(route)
    valid_files=[]
    for file in all_files:
        f=remove_file_extension(file)
        if f.startswith(prefix) and f.endswith(suffix):
            valid_files.append(file)
    return valid_files

#return pairs of freq_files and sp_files which share the same name (which is also included).
# It takes arrays with the freq and sp files that can contain the whole path and extension, 
# and the prefix and sufix for freq and sp files 
def pair_up_freq_sp(freq_files,sp_files,freq_prefix,freq_suffix,sp_prefix,sp_suffix):
    pairs=[]
    i=0
    while i<len(freq_files):
        freq_file_without_ext=remove_file_extension(split_route_and_file(freq_files[i])[1])

        # if the freq and sp files are in the same folder, freq_files and sp_files arrays will be identical
        # the first condition checks if an entry in freq_files array could be a freq file; the second, that it is not an sp file        
        if (freq_file_without_ext.startswith(freq_prefix) and freq_file_without_ext.endswith(freq_suffix)):
            if not(freq_file_without_ext.startswith(sp_prefix) and freq_file_without_ext.endswith(sp_suffix)):

                fr_f=freq_file_without_ext.replace(freq_prefix,"").replace(freq_suffix,"")
                j=0
                while j<len(sp_files):
                    sp_file_without_ext=remove_file_extension(split_route_and_file(sp_files[j])[1])

                    # if the freq and sp files are in the same folder, this checks that this file was not identified as a freq file
                    # it uses freq_files and sp_files arrays (with route and extension) to deal with the case in which freq and sp files have 
                    # the same names but are in different folders
                    if freq_files[i]!=sp_files[j]:
                        sp_f=sp_file_without_ext.replace(sp_prefix,"").replace(sp_suffix,"")
                        if fr_f==sp_f:
                            pairs.append([freq_files[i],sp_files[j],fr_f])
                            break
                    j=j+1
        i=i+1
    return pairs

#given a file with frequency information and a file with single-point energy information, 
# searches in their directories for other files, pairing them up, and finding the common name of the two files
# the one that should be called!!!
#####example:
#pairs= list_pair_of_files("/Users/luissimon/calculos/list-mk/Aldolica/2Si/list-ald-2Si-S-Ph-Sistrans-gau2-.out","/Users/luissimon/calculos/list-mk/Aldolica/2Si/sp2/list-ald-2Si-S-Ph-Sistrans-gau2--sp2.out")

def list_pair_of_files(freq_file,sp_file):
    #delete:
    #freq_file=os.path.expanduser(freq_file)
    #sp_file=os.path.expanduser(sp_file)

    if freq_file!=sp_file:
        freq_route=split_route_and_file(freq_file)[0]
        sp_route=split_route_and_file(sp_file)[0]

        freq_files=os.listdir(freq_route)
        sp_files=os.listdir(sp_route)

        #add the route to each file 


        freq_files_and_route=[]
        sp_files_and_route=[]
        for f in freq_files:
            if os.path.isfile(freq_route+f) and (not f.startswith(".")):
                freq_files_and_route.append(freq_route+f)
        for f in sp_files:
            if os.path.isfile(sp_route+f) and (not f.startswith(".")):
                sp_files_and_route.append(sp_route+f)   

        
        # now find the prefix and suffix for freq and sp calculations
        [[freq_prefix,freq_suffix],[sp_prefix,sp_suffix]]=get_prefix_and_suffix(split_route_and_file(freq_file)[1],split_route_and_file(sp_file)[1])
        # make the magic
        pairs=pair_up_freq_sp(freq_files_and_route,sp_files_and_route,freq_prefix,freq_suffix,sp_prefix,sp_suffix)

    else:
        route=split_route_and_file(freq_file)[0]
        files=os.listdir(freq_route)
        pairs=[]
        for f in files:
            if os.path.isfile(route+f) and (not f.startswith(".")):
                pairs.append([route+f,route+f,remove_file_extension(f)])

    return pairs

######################End of functions for working with files and directories#########################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# Functions for for printing output in the shell:             -A progress bar.                       #
#                                                             -Function to show pairs of files.      #
#                                                                                                    #
######################################################################################################
def progress_bar(counter,max_counter,text):
    if counter==-1:
        sys.stdout.write("progress" + ": " + u"\u2591"*40 + " "*60)
        sys.stdout.flush()

    else:
        progress=int(40*(float(counter)/float(max_counter)))
        sys.stdout.write("\rprogress" + ": " + u"\u2593"*progress + u"\u2591"*(40-progress)+" " + str(counter)+"/"+str(max_counter)+ " "+text)
        sys.stdout.flush()


def print_pairs (f,e,n):
    w=60+len(n)
    if len(f)>(w-30): f="..."+f[-(w-35):]
    if len(e)>(w-30): e="..."+e[-(w-35):]
    print u"\u2582"*(w)
    print u"\u2594"*30+n+u"\u2594"*30
    print "freq. will be read from:      "+f
    print "energy will be read from:     "+e
    print u"\u2594"*(w)

#########################End of functions for printing output in the shell############################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


######################################################################################################
#                                                                                                    #
# to write some random text in the place where the title should be...                                #
#                                                                                                    #
######################################################################################################
def quotes():
    quotes=[]
    quotes.append("We realize the importance of our voices only when we are silenced.")
    quotes.append("One child, one teacher, one book, one pen can change the world.")
    quotes.append("When the whole world is silent, even one voice becomes powerful.")
    quotes.append("I raise up my voice-not so I can shout but so that those without a voice can be heard...we cannot succeed when half of us are held back.")
    quotes.append("If one man can destroy everything, why can't one girl change it?")
    quotes.append("Let us pick up our books and our pens, they are the most powerful weapons.")
    quotes.append("The extremists are afraid of books and pens, the power of education frightens them. they are afraid of women.")
    quotes.append("Once I had asked God for one or two extra inches in height, but instead he made me as tall as the sky, so high that I could not measure myself.")
    quotes.append("My mother always told me, 'hide your face people are looking at you'. I would reply, 'it does not matter; I am also looking at them'.")
    quotes.append("With guns you can kill terrorists, with education you can kill terrorism.")
    quotes.append("I don't want revenge on the Taliban, I want education for sons and daughters of the Taliban.")
    quotes.append("Life isn't just about taking in oxygen and giving out carbon dioxide.")
    quotes.append("I don't want to be thought of as the 'girl who was shot by the Taliban' but the 'girl who fought for education.'")
    quotes.append("We liked to be known as the clever girls. When we decorated our hands with henna for holidays and weddings, we drew calculus and chemical formulae instead of flowers and butterflies.")
    quotes.append("We were scared, but our fear was not as strong as our courage.")
    quotes.append("There are two powers in the world; one is the sword and the other is the pen. There is a third power stronger than both, that of women.")

    return "'"+quotes[random.randint(0,15)]+"'\n (Malala Yousafzai).\n"
############################End of functions for writting random text#################################
###############ENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDENDNDENDEND#############


        
if __name__ == "__main__":

    #default values:
    pml_file="/Users/luissimon/bolillas.pml";
    doc_file_name="SI.doc"
    sp_file=""
    freq_file=""
    keep_files="n"
    include_picture="n"
    optz_pict=""
    picture_x_resolution=800
    picture_y_resolution=800
    pict_size=7
    quiet="n"

    if len(sys.argv) > 1:
        i=0
        while i< len(sys.argv):
            if sys.argv[i]=="-freq_file":
                freq_file=os.path.expanduser(sys.argv[i+1])
                i=i+1
            if sys.argv[i]=="-energy_file":
                sp_file=os.path.expanduser(sys.argv[i+1])
                i=i+1
            if sys.argv[i]=="-pml_file":
                pml_file=os.path.expanduser(sys.argv[i+1])
                i=i+1
            if sys.argv[i]=="-doc_file":
                doc_file_name=os.path.expanduser(sys.argv[i+1])
                i=i+1
            if sys.argv[i]=="-keep_files":
                keep_files="y"
            if sys.argv[i]=="-incl_pict":
                include_picture="y"
            if sys.argv[i]=="-optz_pict":
                optz_pict=sys.argv[i+1]
                i=i+1
            if sys.argv[i]=="-pict_res_x":
                picture_x_resolution=int(sys.argv[i+1])
                i=i+1
            if sys.argv[i]=="-pict_res_y":
                picture_y_resolution=int(sys.argv[i+1])
                i=i+1  
            if sys.argv[i]=="-pict_size":
                pict_size=int(sys.argv[i+1])
                i=i+1 
            if sys.argv[i]=="-q":
                quiet="y"
                i=i+1



            i=i+1
    
    if sp_file=="" and freq_file=="":
        print ("must provide at least the name and route for a file containing the energy or the frequency information")
        sys.exit()
    if sp_file=="" and freq_file!="":
        sp_file=freq_file
    if freq_file=="" and sp_file!="":
        freq_file=sp_file
    
    if (pml_file=="" or os.path.isfile(pml_file)==False) and include_picture=="y":
        print ("WARNING: pictures of the structures will be rendered using default pymol settings...")

    if (doc_file_name.endswith(".doc")==False): doc_file_name=doc_file_name+".doc"
    


    # first let's get the list of files on which we will work....
    print freq_file
    print sp_file
    pairs=list_pair_of_files(freq_file,sp_file)

    if quiet!="y":
        right=""
        for p in pairs:
            print_pairs(p[0],p[1],p[2])
        right = raw_input ("Is everything correct? ")
        if right!="y" and right!="yes" and right!="si" and right!="s":
            print "Sorry, I did my best. Rename the files, put them in folders, delete files that are not neccessary.... and try again!"
            sys.exit()

    # if agree, create the docs.Document object to write...
    doc=create_and_format_document(quotes())

    # and from that list....
    counter=0
    max_counter=len(pairs)
    progress_bar(counter,max_counter,"")
  

    for structure_files in pairs:

        progress_bar(counter,max_counter," structures completed; reading.                             " )
        name= structure_files[2]
        progress_bar(counter,max_counter," structures completed; reading..                            " )
        energy= read_text_energy(structure_files[1])
        progress_bar(counter,max_counter," structures completed; reading...                           " )
        freq= read_img_freqs(structure_files[0])
        progress_bar(counter,max_counter," structures completed; reading....                          " )
        molecule=read_cartsandlayer_from_g_out(structure_files[0])
        cartesians=prepare_cart(molecule)
        progress_bar(counter,max_counter," structures completed; reading.....                         " )
        png_file=""

        if include_picture=="y":

            if optz_pict!="n" and optz_pict!="" and (optz_pict.find("prevent_eclipsing")>-1 or optz_pict.find("eclipse")>-1 or optz_pict.find("span")>-1):
                progress_bar(counter,max_counter," structures completed; rotating...                          " )
                best_orientation=get_best_orientation(molecule,optz_pict)
                molecule_oriented=rotate_molecule(molecule,best_orientation)
            else: molecule_oriented=molecule
                #molecule_oriented=rotate_molecule(molecule,[0,0,0]) ...needed?

            progress_bar(counter,max_counter," structures completed; generating pdb...                    " )            
            molecule_to_pdb(molecule_oriented,name+".pdb")
            progress_bar(counter,max_counter," structures completed; rendering image...                   " ) 
            render_image(name+".pdb",pml_file)
            png_file=name+".png"
        
        if include_picture=="n": png_file=""

        progress_bar(counter,max_counter," structures completed; writting in document...              " )   
    
        write_in_document(doc,name,energy,freq,cartesians,png_file,pict_size)

        #cleaning up... if required
        if keep_files=="n" and include_picture=="y":
            progress_bar(counter,max_counter," structures completed; cleaning up...                      " )
            os.remove(name+".pdb")
            os.remove(name+".png")

        counter=counter+1
        progress_bar(counter,max_counter," structures completed                                          " )




    progress_bar(counter,max_counter," structures completed; finished: "+doc_file_name+ "           \n")    
    # will it be good to save the document only at the end?


    save_document(doc,doc_file_name)
